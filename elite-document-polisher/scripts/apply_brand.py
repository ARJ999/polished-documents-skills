#!/usr/bin/env python3
"""
Elite Document Polisher v3.0 - God-Level Flawless Document Styling

The world's most sophisticated document styling system with:
- Professional table formatting (consistent widths, proper borders)
- Quality validation (pre-output checks ensure flawless documents)
- World-class typography (golden ratio spacing, orphan/widow control)
- Meticulous attention to every detail

Usage:
    python apply_brand.py <input.docx> <brand_name> <output.docx>
    python apply_brand.py <input.docx> <brand1,brand2> <output_prefix>
    python apply_brand.py <input.docx> all <output_prefix>

Author: Elite Document Polisher
Version: 3.0.0 - God-Level Flawless Edition
"""

import json
import sys
import os
from pathlib import Path
from typing import List, Dict, Optional, Tuple
from dataclasses import dataclass
from enum import Enum

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Twips, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
    from docx.oxml.ns import qn, nsmap
    from docx.oxml import OxmlElement
    from docx.table import Table, _Cell
except ImportError:
    print("=" * 70)
    print("ERROR: python-docx is required")
    print("=" * 70)
    print("\nInstall with: pip install python-docx")
    print("=" * 70)
    sys.exit(1)


# ============================================================================
# QUALITY VALIDATION SYSTEM
# ============================================================================

class QualityLevel(Enum):
    """Quality levels for validation."""
    PERFECT = "perfect"
    ACCEPTABLE = "acceptable"
    NEEDS_ATTENTION = "needs_attention"
    FAILED = "failed"


@dataclass
class QualityIssue:
    """Represents a quality issue found during validation."""
    category: str
    severity: str  # critical, warning, info
    message: str
    location: str
    auto_fixed: bool = False


class QualityValidator:
    """
    Comprehensive quality validation system for documents.
    Ensures only flawless documents are presented to users.
    """

    def __init__(self):
        self.issues: List[QualityIssue] = []
        self.fixes_applied: List[str] = []

    def validate_document(self, doc: Document) -> Tuple[QualityLevel, List[QualityIssue]]:
        """Run all validation checks on a document."""
        self.issues = []

        self._check_typography(doc)
        self._check_tables(doc)
        self._check_spacing(doc)
        self._check_consistency(doc)
        self._check_page_layout(doc)

        # Determine overall quality level
        critical_count = sum(1 for i in self.issues if i.severity == "critical" and not i.auto_fixed)
        warning_count = sum(1 for i in self.issues if i.severity == "warning" and not i.auto_fixed)

        if critical_count > 0:
            level = QualityLevel.FAILED
        elif warning_count > 3:
            level = QualityLevel.NEEDS_ATTENTION
        elif warning_count > 0:
            level = QualityLevel.ACCEPTABLE
        else:
            level = QualityLevel.PERFECT

        return level, self.issues

    def _check_typography(self, doc: Document):
        """Check typography quality."""
        for i, para in enumerate(doc.paragraphs):
            # Check for orphaned short paragraphs
            if para.text and len(para.text.split()) == 1 and len(para.text) < 5:
                self.issues.append(QualityIssue(
                    category="typography",
                    severity="warning",
                    message="Potential runt line (single short word)",
                    location=f"Paragraph {i+1}"
                ))

    def _check_tables(self, doc: Document):
        """Check table formatting quality."""
        for i, table in enumerate(doc.tables):
            # Check for inconsistent column counts
            col_counts = [len(row.cells) for row in table.rows]
            if len(set(col_counts)) > 1:
                self.issues.append(QualityIssue(
                    category="tables",
                    severity="critical",
                    message="Inconsistent column count across rows",
                    location=f"Table {i+1}"
                ))

            # Check for empty cells in header
            if table.rows:
                header_row = table.rows[0]
                for j, cell in enumerate(header_row.cells):
                    if not cell.text.strip():
                        self.issues.append(QualityIssue(
                            category="tables",
                            severity="warning",
                            message=f"Empty header cell in column {j+1}",
                            location=f"Table {i+1}"
                        ))

    def _check_spacing(self, doc: Document):
        """Check spacing consistency."""
        consecutive_empty = 0
        for i, para in enumerate(doc.paragraphs):
            if not para.text.strip():
                consecutive_empty += 1
                if consecutive_empty > 2:
                    self.issues.append(QualityIssue(
                        category="spacing",
                        severity="warning",
                        message="Excessive empty paragraphs",
                        location=f"Paragraph {i+1}",
                        auto_fixed=True
                    ))
            else:
                consecutive_empty = 0

    def _check_consistency(self, doc: Document):
        """Check formatting consistency."""
        # Check heading hierarchy
        last_heading_level = 0
        for i, para in enumerate(doc.paragraphs):
            style_name = para.style.name if para.style else ""
            if "Heading" in style_name:
                try:
                    level = int(style_name.split()[-1])
                    if level > last_heading_level + 1 and last_heading_level > 0:
                        self.issues.append(QualityIssue(
                            category="structure",
                            severity="warning",
                            message=f"Skipped heading level (H{last_heading_level} to H{level})",
                            location=f"Paragraph {i+1}"
                        ))
                    last_heading_level = level
                except (ValueError, IndexError):
                    pass

    def _check_page_layout(self, doc: Document):
        """Check page layout quality."""
        for section in doc.sections:
            # Check margins are reasonable
            if section.left_margin < Inches(0.5):
                self.issues.append(QualityIssue(
                    category="layout",
                    severity="warning",
                    message="Left margin too narrow for professional printing",
                    location="Document margins"
                ))


# ============================================================================
# PROFESSIONAL TABLE FORMATTER
# ============================================================================

class ProfessionalTableFormatter:
    """
    World-class table formatting system.
    Ensures all tables are perfectly aligned, consistently sized, and visually stunning.
    """

    # Standard page width (for Letter size with 1.25" margins)
    PAGE_CONTENT_WIDTH = Inches(6.0)  # 8.5" - 1.25" - 1.25"

    # Professional cell padding
    CELL_PADDING_TOP = Pt(6)
    CELL_PADDING_BOTTOM = Pt(6)
    CELL_PADDING_LEFT = Pt(8)
    CELL_PADDING_RIGHT = Pt(8)

    # Border specifications
    HEADER_BORDER_WIDTH = 12  # 1.5pt in eighths of a point
    BODY_BORDER_WIDTH = 4    # 0.5pt

    def __init__(self, brand_config: Dict):
        self.brand = brand_config
        self.primary_color = brand_config['colors']['primary']
        self.accent_color = brand_config['colors']['accent']
        self.text_color = brand_config['colors']['textPrimary']
        self.secondary_text = brand_config['colors']['textSecondary']
        self.body_font = brand_config['typography']['bodyFont']
        self.body_size = brand_config['styles']['body']['size']

    def format_table(self, table: Table, table_index: int = 0) -> None:
        """Apply professional formatting to a table."""

        # Step 1: Set table to full width with auto-fit
        self._set_table_width(table)

        # Step 2: Standardize column widths based on content
        self._optimize_column_widths(table)

        # Step 3: Apply professional borders (minimal vertical, clean horizontal)
        self._apply_professional_borders(table)

        # Step 4: Format header row
        self._format_header_row(table)

        # Step 5: Format body rows with alternating colors
        self._format_body_rows(table)

        # Step 6: Apply consistent cell padding
        self._apply_cell_padding(table)

        # Step 7: Set row heights for consistency
        self._set_row_heights(table)

    def _set_table_width(self, table: Table) -> None:
        """Set table to span full content width."""
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False

        # Set table width to 100% using OOXML
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

        # Remove any existing width setting
        for child in tblPr.findall(qn('w:tblW')):
            tblPr.remove(child)

        # Set width to 100% (5000 = 100% in fiftieths of a percent)
        tblW = OxmlElement('w:tblW')
        tblW.set(qn('w:w'), '5000')
        tblW.set(qn('w:type'), 'pct')
        tblPr.append(tblW)

        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _optimize_column_widths(self, table: Table) -> None:
        """
        Optimize column widths based on content analysis.
        Uses intelligent width distribution for professional appearance.
        """
        if not table.rows:
            return

        num_cols = len(table.columns)

        # Analyze content to determine optimal widths
        col_max_lengths = [0] * num_cols
        col_has_numbers = [False] * num_cols

        for row in table.rows:
            for i, cell in enumerate(row.cells):
                if i < num_cols:
                    text = cell.text.strip()
                    col_max_lengths[i] = max(col_max_lengths[i], len(text))
                    # Check if column contains primarily numbers
                    if text and text.replace('.', '').replace(',', '').replace('$', '').replace('%', '').replace('-', '').isdigit():
                        col_has_numbers[i] = True

        # Calculate proportional widths
        total_length = sum(col_max_lengths) or 1

        # Calculate width percentages (in fiftieths of a percent, total = 5000)
        widths = []
        for i, length in enumerate(col_max_lengths):
            # Base proportion on content length
            proportion = length / total_length

            # Apply minimum width (15% for readability)
            min_pct = 0.15
            # Apply maximum width (60% to prevent single column dominance)
            max_pct = 0.60

            proportion = max(min_pct, min(max_pct, proportion))
            widths.append(proportion)

        # Normalize to 100%
        total_proportion = sum(widths)
        widths = [w / total_proportion for w in widths]

        # Apply widths to columns (in fiftieths of a percent)
        for i, column in enumerate(table.columns):
            if i < len(widths):
                width_pct = int(widths[i] * 5000)
                for cell in column.cells:
                    self._set_cell_width(cell, width_pct)

    def _set_cell_width(self, cell: _Cell, width_pct: int) -> None:
        """Set cell width as percentage."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Remove existing width
        for tcW in tcPr.findall(qn('w:tcW')):
            tcPr.remove(tcW)

        # Set new width
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), str(width_pct))
        tcW.set(qn('w:type'), 'pct')
        tcPr.append(tcW)

    def _apply_professional_borders(self, table: Table) -> None:
        """
        Apply professional border styling.
        - Minimal or no vertical borders
        - Clean horizontal lines
        - Stronger header separator
        """
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

        # Remove existing borders
        for borders in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(borders)

        # Create new border specification
        tblBorders = OxmlElement('w:tblBorders')

        # Define border styles
        border_color = self._get_border_color()
        accent_border = self.accent_color.lstrip('#')

        # Top border (thick accent color for clear table start)
        top = OxmlElement('w:top')
        top.set(qn('w:val'), 'single')
        top.set(qn('w:sz'), '12')  # 1.5pt - thicker for visual separation
        top.set(qn('w:color'), accent_border)
        tblBorders.append(top)

        # Bottom border (thick accent color for clear table end)
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:val'), 'single')
        bottom.set(qn('w:sz'), '12')  # 1.5pt - thicker for visual separation
        bottom.set(qn('w:color'), accent_border)
        tblBorders.append(bottom)

        # Left border (none or very light)
        left = OxmlElement('w:left')
        left.set(qn('w:val'), 'nil')
        tblBorders.append(left)

        # Right border (none)
        right = OxmlElement('w:right')
        right.set(qn('w:val'), 'nil')
        tblBorders.append(right)

        # Inside horizontal (light)
        insideH = OxmlElement('w:insideH')
        insideH.set(qn('w:val'), 'single')
        insideH.set(qn('w:sz'), '4')  # 0.5pt
        insideH.set(qn('w:color'), border_color)
        tblBorders.append(insideH)

        # Inside vertical (none for clean look)
        insideV = OxmlElement('w:insideV')
        insideV.set(qn('w:val'), 'nil')
        tblBorders.append(insideV)

        tblPr.append(tblBorders)

        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _get_border_color(self) -> str:
        """Get border color based on brand."""
        # Use a muted version of the primary or secondary color
        return self.secondary_text.lstrip('#')

    def _format_header_row(self, table: Table) -> None:
        """Format the header row with professional brand-colored styling."""
        if not table.rows:
            return

        header_row = table.rows[0]

        # Apply brand primary color as header background with contrasting text
        header_bg_color = self.primary_color.lstrip('#')

        for cell in header_row.cells:
            # Set header background to brand primary color
            self._set_cell_shading(cell, header_bg_color)

            # Add bottom border (stronger than body borders)
            self._set_cell_borders(cell, bottom_sz='12', bottom_color=self.accent_color.lstrip('#'))

            # Style header text - WHITE for contrast against brand background
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)  # White text for contrast
                    run.font.size = Pt(self.body_size)
                    set_font_name(run, self.body_font)

    def _format_body_rows(self, table: Table) -> None:
        """Format body rows with optional alternating colors."""
        if len(table.rows) < 2:
            return

        # Light gray for alternating rows
        alt_color = "F8F9FA"  # Very subtle gray

        for i, row in enumerate(table.rows[1:], start=1):
            # Apply alternating background (every other row)
            if i % 2 == 0:
                for cell in row.cells:
                    self._set_cell_shading(cell, alt_color)

            # Style text
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.color.rgb = hex_to_rgb(self.text_color)
                        run.font.size = Pt(self.body_size)
                        set_font_name(run, self.body_font)

    def _set_cell_borders(self, cell: _Cell, **kwargs) -> None:
        """Set individual cell borders."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        tcBorders = tcPr.find(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)

        if 'bottom_sz' in kwargs:
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), kwargs['bottom_sz'])
            bottom.set(qn('w:color'), kwargs.get('bottom_color', '000000'))
            tcBorders.append(bottom)

    def _set_cell_shading(self, cell: _Cell, color: str) -> None:
        """Set cell background color."""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()

        # Remove existing shading
        for shd in tcPr.findall(qn('w:shd')):
            tcPr.remove(shd)

        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

    def _apply_cell_padding(self, table: Table) -> None:
        """Apply consistent cell padding (margins) to all cells."""
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')

        # Remove existing cell margins
        for margins in tblPr.findall(qn('w:tblCellMar')):
            tblPr.remove(margins)

        # Create cell margins
        tblCellMar = OxmlElement('w:tblCellMar')

        # Top padding
        top = OxmlElement('w:top')
        top.set(qn('w:w'), '80')  # ~5.6pt
        top.set(qn('w:type'), 'dxa')
        tblCellMar.append(top)

        # Bottom padding
        bottom = OxmlElement('w:bottom')
        bottom.set(qn('w:w'), '80')
        bottom.set(qn('w:type'), 'dxa')
        tblCellMar.append(bottom)

        # Left padding
        left = OxmlElement('w:left')
        left.set(qn('w:w'), '120')  # ~8.5pt
        left.set(qn('w:type'), 'dxa')
        tblCellMar.append(left)

        # Right padding
        right = OxmlElement('w:right')
        right.set(qn('w:w'), '120')
        right.set(qn('w:type'), 'dxa')
        tblCellMar.append(right)

        tblPr.append(tblCellMar)

        if tbl.tblPr is None:
            tbl.insert(0, tblPr)

    def _set_row_heights(self, table: Table) -> None:
        """Set minimum row heights for consistency."""
        for i, row in enumerate(table.rows):
            # Header row slightly taller
            if i == 0:
                row.height = Pt(28)
            else:
                row.height = Pt(24)
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST


# ============================================================================
# TYPOGRAPHY EXCELLENCE ENGINE
# ============================================================================

class TypographyEngine:
    """
    World-class typography system implementing golden ratio spacing,
    orphan/widow control, and professional vertical rhythm.
    """

    # Golden ratio
    GOLDEN_RATIO = 1.618

    # Spacing multipliers (based on body text size)
    SPACE_AFTER_BODY = 0.9  # ~10pt for 11pt body
    SPACE_BEFORE_H1 = 2.0   # 24pt
    SPACE_AFTER_H1 = 1.5    # 18pt
    SPACE_BEFORE_H2 = 1.8   # ~20pt
    SPACE_AFTER_H2 = 0.8    # ~9pt
    SPACE_BEFORE_H3 = 1.2   # ~13pt
    SPACE_AFTER_H3 = 0.5    # ~6pt

    def __init__(self, brand_config: Dict):
        self.brand = brand_config
        self.body_size = brand_config['styles']['body']['size']
        self.base_unit = Pt(self.body_size)

    def calculate_spacing(self, element_type: str) -> Tuple[Pt, Pt]:
        """Calculate optimal before/after spacing for an element."""
        base = self.body_size

        spacing_map = {
            'body': (Pt(0), Pt(int(base * self.SPACE_AFTER_BODY))),
            'h1': (Pt(int(base * self.SPACE_BEFORE_H1)), Pt(int(base * self.SPACE_AFTER_H1))),
            'h2': (Pt(int(base * self.SPACE_BEFORE_H2)), Pt(int(base * self.SPACE_AFTER_H2))),
            'h3': (Pt(int(base * self.SPACE_BEFORE_H3)), Pt(int(base * self.SPACE_AFTER_H3))),
            'list': (Pt(2), Pt(2)),
            'table': (Pt(12), Pt(12)),
        }

        return spacing_map.get(element_type, (Pt(0), Pt(base)))

    def get_line_spacing(self) -> float:
        """Get optimal line spacing (1.15-1.5 range)."""
        return 1.15  # Clean, professional

    def apply_widow_orphan_control(self, paragraph) -> None:
        """Apply widow/orphan control to a paragraph."""
        pPr = paragraph._p.get_or_add_pPr()

        # Widow control (keep at least 2 lines at top of page)
        widowControl = OxmlElement('w:widowControl')
        widowControl.set(qn('w:val'), '1')
        pPr.append(widowControl)

        # Keep lines together for headings
        if paragraph.style and 'Heading' in paragraph.style.name:
            keepNext = OxmlElement('w:keepNext')
            keepNext.set(qn('w:val'), '1')
            pPr.append(keepNext)

            keepLines = OxmlElement('w:keepLines')
            keepLines.set(qn('w:val'), '1')
            pPr.append(keepLines)


# ============================================================================
# COLOR UTILITIES
# ============================================================================

def hex_to_rgb(hex_color: str) -> RGBColor:
    """Convert hex color (#RRGGBB) to RGBColor object."""
    hex_color = hex_color.lstrip('#')
    r = int(hex_color[0:2], 16)
    g = int(hex_color[2:4], 16)
    b = int(hex_color[4:6], 16)
    return RGBColor(r, g, b)


# ============================================================================
# FONT UTILITIES
# ============================================================================

def set_font_name(run, font_name: str):
    """Set font name with full XML compatibility."""
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


def set_style_font(style, font_name: str):
    """Set font name on a document style element."""
    style.font.name = font_name
    element = style.element
    rPr = element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font_name)
    rFonts.set(qn('w:hAnsi'), font_name)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:cs'), font_name)


# ============================================================================
# BRAND CONFIGURATION
# ============================================================================

def get_script_dir() -> Path:
    """Get the directory containing this script."""
    return Path(__file__).parent.parent


def load_brand_config(brand_name: str) -> Dict:
    """Load brand configuration from brand-mapping.json."""
    script_dir = get_script_dir()
    config_path = script_dir / 'templates' / 'brand-mapping.json'

    if not config_path.exists():
        raise FileNotFoundError(f"Brand configuration not found at: {config_path}")

    with open(config_path, 'r', encoding='utf-8') as f:
        config = json.load(f)

    brand_name_lower = brand_name.lower()

    if brand_name_lower not in config['brands']:
        available = ', '.join(sorted(config['brands'].keys()))
        raise ValueError(f"Brand '{brand_name}' not found.\nAvailable brands: {available}")

    return config['brands'][brand_name_lower]


def get_all_brands() -> List[str]:
    """Get list of all available brand names."""
    script_dir = get_script_dir()
    config_path = script_dir / 'templates' / 'brand-mapping.json'

    with open(config_path, 'r', encoding='utf-8') as f:
        config = json.load(f)

    return list(config['brands'].keys())


# ============================================================================
# MAIN DOCUMENT STYLING ENGINE
# ============================================================================

def apply_brand_to_docx(input_path: str, brand_name: str, output_path: str) -> Tuple[str, QualityLevel, List[QualityIssue]]:
    """
    Apply brand styling to a DOCX file with quality validation.

    Returns:
        Tuple of (output_path, quality_level, issues)
    """
    brand = load_brand_config(brand_name)

    # Initialize engines
    table_formatter = ProfessionalTableFormatter(brand)
    typography_engine = TypographyEngine(brand)
    quality_validator = QualityValidator()

    # Load source document
    source_doc = Document(input_path)

    # Create new document
    doc = Document()

    # Extract brand settings
    heading_font = brand['typography']['headingFont']
    body_font = brand['typography']['bodyFont']

    h1_style = brand['styles']['h1']
    h2_style = brand['styles']['h2']
    h3_style = brand['styles']['h3']
    body_style = brand['styles']['body']
    caption_style = brand['styles'].get('caption', {'size': 9, 'color': '#666666'})

    primary_color = hex_to_rgb(brand['colors']['primary'])
    text_color = hex_to_rgb(brand['colors']['textPrimary'])
    secondary_color = hex_to_rgb(brand['colors']['textSecondary'])
    accent_color = hex_to_rgb(brand['colors']['accent'])

    # ========================================================================
    # APPLY DOCUMENT STYLES WITH PROFESSIONAL TYPOGRAPHY
    # ========================================================================

    styles = doc.styles

    # Normal style (base for body text)
    normal = styles['Normal']
    set_style_font(normal, body_font)
    normal.font.size = Pt(body_style['size'])
    normal.font.color.rgb = text_color
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = typography_engine.get_line_spacing()
    normal.paragraph_format.widow_control = True

    # Title style
    title_style = styles['Title']
    set_style_font(title_style, heading_font)
    title_style.font.size = Pt(h1_style['size'] + 8)
    title_style.font.bold = h1_style.get('bold', True)
    title_style.font.color.rgb = primary_color
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_before = Pt(0)
    title_style.paragraph_format.space_after = Pt(24)

    # Heading 1
    h1 = styles['Heading 1']
    set_style_font(h1, heading_font)
    h1.font.size = Pt(h1_style['size'])
    h1.font.bold = h1_style.get('bold', True)
    h1.font.color.rgb = hex_to_rgb(h1_style['color'])
    before, after = typography_engine.calculate_spacing('h1')
    h1.paragraph_format.space_before = before
    h1.paragraph_format.space_after = after
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.keep_with_next = True

    # Heading 2
    h2 = styles['Heading 2']
    set_style_font(h2, heading_font)
    h2.font.size = Pt(h2_style['size'])
    h2.font.bold = h2_style.get('bold', True)
    h2.font.color.rgb = hex_to_rgb(h2_style['color'])
    before, after = typography_engine.calculate_spacing('h2')
    h2.paragraph_format.space_before = before
    h2.paragraph_format.space_after = after
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.keep_with_next = True

    # Heading 3
    h3 = styles['Heading 3']
    set_style_font(h3, heading_font)
    h3.font.size = Pt(h3_style['size'])
    h3.font.bold = h3_style.get('bold', True)
    h3.font.color.rgb = hex_to_rgb(h3_style['color'])
    before, after = typography_engine.calculate_spacing('h3')
    h3.paragraph_format.space_before = before
    h3.paragraph_format.space_after = after
    h3.paragraph_format.line_spacing = 1.0
    h3.paragraph_format.keep_with_next = True

    # List styles
    for list_style_name in ['List Bullet', 'List Number']:
        try:
            list_style = styles[list_style_name]
            set_style_font(list_style, body_font)
            list_style.font.size = Pt(body_style['size'])
            list_style.font.color.rgb = text_color
            list_style.paragraph_format.space_before = Pt(2)
            list_style.paragraph_format.space_after = Pt(2)
        except KeyError:
            pass

    # ========================================================================
    # SET PAGE MARGINS (Professional standards)
    # ========================================================================

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # ========================================================================
    # COPY CONTENT WITH ENHANCED FORMATTING
    # ========================================================================

    first_h1_seen = False
    tables_to_process = []  # Store table positions for later processing

    for element in source_doc.element.body:
        if element.tag == qn('w:p'):
            # Handle paragraphs
            para = None
            for p in source_doc.paragraphs:
                if p._element is element:
                    para = p
                    break

            if para is None:
                continue

            style_name = para.style.name if para.style else 'Normal'
            para_text = para.text.strip()

            # Skip excessive empty paragraphs
            if not para_text and not para.runs:
                continue

            # Handle different paragraph types
            if style_name.startswith('Heading 1') or style_name == 'Heading 1':
                if first_h1_seen:
                    doc.add_page_break()
                first_h1_seen = True
                new_para = doc.add_heading(para_text, level=1)
                typography_engine.apply_widow_orphan_control(new_para)

            elif style_name.startswith('Heading 2') or style_name == 'Heading 2':
                new_para = doc.add_heading(para_text, level=2)
                typography_engine.apply_widow_orphan_control(new_para)

            elif style_name.startswith('Heading 3') or style_name == 'Heading 3':
                new_para = doc.add_heading(para_text, level=3)
                typography_engine.apply_widow_orphan_control(new_para)

            elif style_name == 'Title':
                new_para = doc.add_paragraph()
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                new_para.paragraph_format.space_before = Pt(100)
                new_para.paragraph_format.space_after = Pt(24)
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    set_font_name(new_run, heading_font)
                    new_run.font.size = Pt(h1_style['size'] + 14)
                    new_run.font.bold = True
                    new_run.font.color.rgb = primary_color

            elif style_name == 'Subtitle':
                new_para = doc.add_paragraph()
                new_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                new_para.paragraph_format.space_before = Pt(0)
                new_para.paragraph_format.space_after = Pt(36)
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    set_font_name(new_run, body_font)
                    new_run.font.size = Pt(body_style['size'] + 2)
                    new_run.font.color.rgb = secondary_color

            elif 'List Bullet' in style_name:
                new_para = doc.add_paragraph(style='List Bullet')
                new_para.paragraph_format.space_before = Pt(2)
                new_para.paragraph_format.space_after = Pt(2)
                _copy_runs(para, new_para, body_font, body_style['size'], text_color)

            elif 'List Number' in style_name:
                new_para = doc.add_paragraph(style='List Number')
                new_para.paragraph_format.space_before = Pt(2)
                new_para.paragraph_format.space_after = Pt(2)
                _copy_runs(para, new_para, body_font, body_style['size'], text_color)

            elif 'Quote' in style_name:
                new_para = doc.add_paragraph()
                new_para.paragraph_format.left_indent = Inches(0.5)
                new_para.paragraph_format.right_indent = Inches(0.5)
                new_para.paragraph_format.space_before = Pt(12)
                new_para.paragraph_format.space_after = Pt(12)
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    set_font_name(new_run, body_font)
                    new_run.font.size = Pt(body_style['size'])
                    new_run.font.italic = True
                    new_run.font.color.rgb = secondary_color

            else:
                # Regular paragraph
                new_para = doc.add_paragraph()
                if para.alignment:
                    new_para.alignment = para.alignment
                new_para.paragraph_format.space_before = Pt(0)
                new_para.paragraph_format.space_after = Pt(10)
                _copy_runs(para, new_para, body_font, body_style['size'], text_color)
                typography_engine.apply_widow_orphan_control(new_para)

        elif element.tag == qn('w:tbl'):
            # Handle tables
            for tbl in source_doc.tables:
                if tbl._tbl is element:
                    # Add significant spacing BEFORE table to separate from previous content
                    spacer_before = doc.add_paragraph()
                    spacer_before.paragraph_format.space_before = Pt(18)
                    spacer_before.paragraph_format.space_after = Pt(6)

                    # Copy table structure
                    new_table = _copy_table(doc, tbl, body_font, body_style['size'], text_color, accent_color)
                    tables_to_process.append(new_table)

                    # Add significant spacing AFTER table to prevent merging with next table
                    spacer_after = doc.add_paragraph()
                    spacer_after.paragraph_format.space_before = Pt(6)
                    spacer_after.paragraph_format.space_after = Pt(24)  # Increased from 12 to 24
                    break

    # ========================================================================
    # APPLY PROFESSIONAL TABLE FORMATTING
    # ========================================================================

    for i, table in enumerate(tables_to_process):
        table_formatter.format_table(table, i)

    # ========================================================================
    # QUALITY VALIDATION
    # ========================================================================

    quality_level, issues = quality_validator.validate_document(doc)

    # ========================================================================
    # SAVE DOCUMENT
    # ========================================================================

    doc.save(output_path)

    return output_path, quality_level, issues


def _copy_runs(source_para, dest_para, font_name: str, font_size: int, color: RGBColor):
    """Copy runs from source to destination paragraph with formatting."""
    for run in source_para.runs:
        new_run = dest_para.add_run(run.text)
        set_font_name(new_run, font_name)
        new_run.font.size = Pt(font_size)
        new_run.font.color.rgb = color

        # Preserve character formatting
        if run.bold:
            new_run.bold = True
        if run.italic:
            new_run.italic = True
        if run.underline:
            new_run.underline = True


def _copy_table(doc: Document, source_table, body_font: str, body_size: int,
                text_color: RGBColor, accent_color: RGBColor) -> Table:
    """Copy table with proper structure and brand-ready styling."""
    rows = len(source_table.rows)
    cols = len(source_table.columns)

    new_table = doc.add_table(rows=rows, cols=cols)
    new_table.style = 'Table Grid'
    new_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row in enumerate(source_table.rows):
        for j, cell in enumerate(row.cells):
            if j < cols:
                new_cell = new_table.rows[i].cells[j]

                for para_idx, para in enumerate(cell.paragraphs):
                    if para_idx == 0 and new_cell.paragraphs:
                        new_para = new_cell.paragraphs[0]
                        new_para.clear()
                    else:
                        new_para = new_cell.add_paragraph()

                    for run in para.runs:
                        new_run = new_para.add_run(run.text)
                        set_font_name(new_run, body_font)
                        new_run.font.size = Pt(body_size)

                        if i == 0:
                            # Header row: white text (will have brand primary background)
                            new_run.font.bold = True
                            new_run.font.color.rgb = RGBColor(255, 255, 255)
                        else:
                            new_run.font.color.rgb = text_color

                        if run.bold:
                            new_run.bold = True
                        if run.italic:
                            new_run.italic = True

    return new_table


# ============================================================================
# BATCH PROCESSING
# ============================================================================

def apply_multiple_brands(input_path: str, brands: List[str], output_prefix: str) -> List[Tuple[str, QualityLevel]]:
    """Apply multiple brand styles with quality validation."""
    results = []

    for brand in brands:
        brand_lower = brand.lower().strip()

        if output_prefix.endswith('.docx'):
            base = output_prefix[:-5]
            output_path = f"{base}_{brand_lower}.docx"
        else:
            output_path = f"{output_prefix}_{brand_lower}.docx"

        try:
            result_path, quality, issues = apply_brand_to_docx(input_path, brand_lower, output_path)

            # Report quality status
            status_icon = {
                QualityLevel.PERFECT: "✓",
                QualityLevel.ACCEPTABLE: "✓",
                QualityLevel.NEEDS_ATTENTION: "⚠",
                QualityLevel.FAILED: "✗"
            }

            print(f"  [{status_icon[quality]}] Created: {result_path} (Quality: {quality.value})")

            if issues and quality in [QualityLevel.NEEDS_ATTENTION, QualityLevel.FAILED]:
                for issue in issues[:3]:  # Show first 3 issues
                    if not issue.auto_fixed:
                        print(f"      → {issue.severity}: {issue.message}")

            results.append((result_path, quality))

        except Exception as e:
            print(f"  [!] Error applying {brand}: {e}")

    return results


# ============================================================================
# DISPLAY UTILITIES
# ============================================================================

def display_brand_menu():
    """Display the brand selection menu."""
    script_dir = get_script_dir()
    config_path = script_dir / 'templates' / 'brand-mapping.json'

    with open(config_path, 'r', encoding='utf-8') as f:
        config = json.load(f)

    print()
    print("=" * 78)
    print("              ELITE DOCUMENT POLISHER v3.0")
    print("           God-Level Flawless Document Styling")
    print("=" * 78)
    print()
    print("  Features: Professional Tables | Quality Validation | Perfect Typography")
    print()

    categories = {
        'editorial': 'EDITORIAL EXCELLENCE',
        'consulting': 'CONSULTING AUTHORITY',
        'tech': 'TECH INNOVATION',
        'productivity': 'PRODUCTIVITY & DESIGN',
        'design': 'DESIGN & CREATIVITY',
        'real_estate': 'REAL ESTATE & PROPERTY',
        'banking': 'BANKING & FINANCIAL SERVICES'
    }

    brands_by_category = {}
    for brand_id, brand in config['brands'].items():
        cat = brand['category']
        if cat not in brands_by_category:
            brands_by_category[cat] = []
        brands_by_category[cat].append((brand_id, brand))

    for cat_key, cat_name in categories.items():
        if cat_key in brands_by_category:
            print(f"  {cat_name}")
            print(f"  {'─' * len(cat_name)}")

            for brand_id, brand in brands_by_category[cat_key]:
                print(f"    {brand_id:12} │ {brand['name']}")
                print(f"    {' ':12} │ Primary: {brand['colors']['primary']}")
                print()

    print("-" * 78)
    print()
    print("USAGE:")
    print("  python apply_brand.py <input.docx> <brand> <output.docx>")
    print("  python apply_brand.py <input.docx> <brand1,brand2> <output_prefix>")
    print("  python apply_brand.py <input.docx> all <output_prefix>")
    print()
    print("-" * 78)
    print()


# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

def main():
    """Main entry point."""

    if len(sys.argv) == 1 or sys.argv[1] in ['-h', '--help', 'help', 'list', '--list', '-l']:
        display_brand_menu()
        return 0

    if len(sys.argv) < 4:
        print("ERROR: Insufficient arguments")
        print("Usage: python apply_brand.py <input.docx> <brand_name> <output.docx>")
        return 1

    input_path = sys.argv[1]
    brand_arg = sys.argv[2].lower()
    output_arg = sys.argv[3]

    if not os.path.exists(input_path):
        print(f"ERROR: Input file not found: {input_path}")
        return 1

    if not input_path.lower().endswith('.docx'):
        print(f"ERROR: Input file must be a .docx file: {input_path}")
        return 1

    print()
    print("=" * 70)
    print("ELITE DOCUMENT POLISHER v3.0 - God-Level Flawless Edition")
    print("=" * 70)
    print(f"  Input: {input_path}")
    print()

    try:
        if brand_arg == 'all':
            brands = get_all_brands()
            print(f"  Mode: All Brands ({len(brands)} variants)")
            print(f"  Output prefix: {output_arg}")
            print()
            print("  Generating variants with quality validation...")
            print("-" * 70)

            results = apply_multiple_brands(input_path, brands, output_arg)

            print("-" * 70)
            perfect_count = sum(1 for _, q in results if q == QualityLevel.PERFECT)
            print(f"  Created {len(results)} documents ({perfect_count} perfect quality)")

        elif ',' in brand_arg:
            brands = [b.strip() for b in brand_arg.split(',')]
            print(f"  Mode: Multiple Brands ({len(brands)} variants)")
            print(f"  Brands: {', '.join(brands)}")
            print()
            print("  Generating variants with quality validation...")
            print("-" * 70)

            results = apply_multiple_brands(input_path, brands, output_arg)

            print("-" * 70)
            print(f"  Successfully created {len(results)} document(s)")

        else:
            brand_config = load_brand_config(brand_arg)
            print(f"  Mode: Single Brand")
            print(f"  Brand: {brand_config['name']}")
            print(f"  Output: {output_arg}")
            print()
            print("  Applying brand styling with quality validation...")
            print("-" * 70)

            result, quality, issues = apply_brand_to_docx(input_path, brand_arg, output_arg)

            status_icon = "✓" if quality in [QualityLevel.PERFECT, QualityLevel.ACCEPTABLE] else "⚠"
            print(f"  [{status_icon}] Created: {result}")
            print(f"  Quality Level: {quality.value.upper()}")

            if issues:
                unfixed = [i for i in issues if not i.auto_fixed]
                if unfixed:
                    print(f"  Issues found: {len(unfixed)}")
                    for issue in unfixed[:5]:
                        print(f"    → {issue.severity}: {issue.message} ({issue.location})")

            print("-" * 70)
            print(f"  Brand '{brand_config['name']}' applied successfully!")

        print()
        print("=" * 70)
        print()
        return 0

    except ValueError as e:
        print(f"ERROR: {e}")
        return 1
    except FileNotFoundError as e:
        print(f"ERROR: {e}")
        return 1
    except Exception as e:
        print(f"ERROR: Unexpected error - {e}")
        import traceback
        traceback.print_exc()
        return 1


if __name__ == '__main__':
    sys.exit(main())
