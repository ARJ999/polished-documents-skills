#!/usr/bin/env python3
"""
Enhanced Sample Document Generator for Elite Document Polisher v3.0

Creates a comprehensive sample document with multiple tables of varying
column counts to test the professional table formatting system.

Usage:
    python create_sample.py

Output:
    sample_report.docx
"""

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    exit(1)


def create_sample_document():
    """Create a comprehensive sample document with multiple table types."""

    doc = Document()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Data Platform Implementation Strategy")
    run.bold = True
    title.style = 'Title'

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Enterprise Analytics Modernization Initiative")
    subtitle.style = 'Subtitle'

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)

    doc.add_paragraph(
        "This document outlines our comprehensive strategy for implementing a modern "
        "data platform that will transform how our organization leverages data for "
        "decision-making. The initiative encompasses data lakehouse architecture, "
        "AI/ML capabilities, and enterprise-wide analytics democratization."
    )

    # Platform Capabilities - 2 Column Table
    doc.add_heading("Platform Capabilities", level=2)

    doc.add_paragraph(
        "The following table summarizes the core capabilities of our proposed "
        "data platform and their direct business impact:"
    )

    # TABLE 1: 2 columns (like user's first table)
    table1 = doc.add_table(rows=4, cols=2)
    table1.style = 'Table Grid'

    # Header
    table1.rows[0].cells[0].text = "Capability"
    table1.rows[0].cells[1].text = "Business Impact"

    # Data rows
    table1.rows[1].cells[0].text = "Real-time data ingestion"
    table1.rows[1].cells[1].text = "Instant insights on sales pipeline and customer behavior"

    table1.rows[2].cells[0].text = "Semantic layer abstraction"
    table1.rows[2].cells[1].text = "Business users can query data using natural language"

    table1.rows[3].cells[0].text = "Built-in ML/AI capabilities"
    table1.rows[3].cells[1].text = "Predictive models without data movement"

    # Implementation Roadmap - 3 Column Table
    doc.add_heading("Implementation Roadmap", level=2)

    doc.add_paragraph(
        "Our phased approach ensures minimal disruption while delivering "
        "incremental value throughout the implementation journey:"
    )

    # TABLE 2: 3 columns (like user's second table)
    table2 = doc.add_table(rows=5, cols=3)
    table2.style = 'Table Grid'

    # Header
    table2.rows[0].cells[0].text = "Phase"
    table2.rows[0].cells[1].text = "Timeline"
    table2.rows[0].cells[2].text = "Key Deliverables"

    # Data rows
    table2.rows[1].cells[0].text = "Foundation"
    table2.rows[1].cells[1].text = "Q1 2026"
    table2.rows[1].cells[2].text = "Data Lakehouse deployment, integration pipelines"

    table2.rows[2].cells[0].text = "Intelligence"
    table2.rows[2].cells[1].text = "Q2 2026"
    table2.rows[2].cells[2].text = "Cortex AI analytics, semantic layer, natural language queries"

    table2.rows[3].cells[0].text = "Optimization"
    table2.rows[3].cells[1].text = "Q3 2026"
    table2.rows[3].cells[2].text = "Predictive models, automated workflows, decision engines"

    table2.rows[4].cells[0].text = "Scale"
    table2.rows[4].cells[1].text = "Q4 2026"
    table2.rows[4].cells[2].text = "Enterprise rollout, continuous improvement, advanced AI use cases"

    # Investment Summary - 4 Column Table
    doc.add_heading("Investment Summary", level=2)

    doc.add_paragraph(
        "The following table provides a breakdown of investment requirements "
        "across different categories:"
    )

    # TABLE 3: 4 columns
    table3 = doc.add_table(rows=5, cols=4)
    table3.style = 'Table Grid'

    # Header
    table3.rows[0].cells[0].text = "Category"
    table3.rows[0].cells[1].text = "Year 1"
    table3.rows[0].cells[2].text = "Year 2"
    table3.rows[0].cells[3].text = "Year 3"

    # Data rows
    table3.rows[1].cells[0].text = "Platform Licensing"
    table3.rows[1].cells[1].text = "$450,000"
    table3.rows[1].cells[2].text = "$425,000"
    table3.rows[1].cells[3].text = "$400,000"

    table3.rows[2].cells[0].text = "Implementation Services"
    table3.rows[2].cells[1].text = "$280,000"
    table3.rows[2].cells[2].text = "$120,000"
    table3.rows[2].cells[3].text = "$80,000"

    table3.rows[3].cells[0].text = "Training & Change Management"
    table3.rows[3].cells[1].text = "$95,000"
    table3.rows[3].cells[2].text = "$45,000"
    table3.rows[3].cells[3].text = "$25,000"

    table3.rows[4].cells[0].text = "Total Investment"
    table3.rows[4].cells[1].text = "$825,000"
    table3.rows[4].cells[2].text = "$590,000"
    table3.rows[4].cells[3].text = "$505,000"

    # Key Benefits
    doc.add_heading("Key Benefits", level=2)

    doc.add_paragraph("The platform will deliver the following strategic benefits:")

    doc.add_paragraph("Unified data foundation eliminating silos across departments", style='List Bullet')
    doc.add_paragraph("Self-service analytics empowering business users", style='List Bullet')
    doc.add_paragraph("AI-powered insights driving proactive decision-making", style='List Bullet')
    doc.add_paragraph("Reduced time-to-insight from weeks to minutes", style='List Bullet')
    doc.add_paragraph("Scalable architecture supporting future growth", style='List Bullet')

    # Risk Assessment
    doc.add_heading("Risk Assessment", level=2)

    doc.add_paragraph(
        "We have identified and developed mitigation strategies for the "
        "following key risks:"
    )

    # Risk table - 3 columns
    table4 = doc.add_table(rows=4, cols=3)
    table4.style = 'Table Grid'

    table4.rows[0].cells[0].text = "Risk"
    table4.rows[0].cells[1].text = "Impact"
    table4.rows[0].cells[2].text = "Mitigation"

    table4.rows[1].cells[0].text = "Data quality issues"
    table4.rows[1].cells[1].text = "High"
    table4.rows[1].cells[2].text = "Automated data quality monitoring and remediation"

    table4.rows[2].cells[0].text = "User adoption resistance"
    table4.rows[2].cells[1].text = "Medium"
    table4.rows[2].cells[2].text = "Comprehensive training program and change champions"

    table4.rows[3].cells[0].text = "Integration complexity"
    table4.rows[3].cells[1].text = "Medium"
    table4.rows[3].cells[2].text = "Phased approach with proven integration patterns"

    # Success Metrics
    doc.add_heading("Success Metrics", level=2)

    doc.add_paragraph(
        "We will measure success through the following key performance indicators:"
    )

    doc.add_paragraph("Data availability: 99.9% uptime for critical data assets", style='List Number')
    doc.add_paragraph("User adoption: 80% of target users actively using the platform", style='List Number')
    doc.add_paragraph("Time-to-insight: 75% reduction in average reporting time", style='List Number')
    doc.add_paragraph("Data quality: 95% accuracy score across all data domains", style='List Number')

    # Next Steps
    doc.add_heading("Next Steps", level=1)

    doc.add_paragraph(
        "To move forward with this initiative, we recommend the following "
        "immediate actions:"
    )

    doc.add_paragraph("Secure executive sponsorship and budget approval", style='List Number')
    doc.add_paragraph("Form cross-functional steering committee", style='List Number')
    doc.add_paragraph("Finalize vendor selection and contract negotiations", style='List Number')
    doc.add_paragraph("Initiate Phase 1 planning and resource allocation", style='List Number')
    doc.add_paragraph("Establish governance framework and data stewardship model", style='List Number')

    # Conclusion
    doc.add_heading("Conclusion", level=1)

    doc.add_paragraph(
        "This data platform initiative represents a transformational opportunity "
        "for our organization. By investing in modern data infrastructure and "
        "AI capabilities, we will establish a competitive advantage that enables "
        "faster, more informed decision-making across all business functions."
    )

    para = doc.add_paragraph()
    para.add_run("We are confident that this investment will deliver ").bold = False
    para.add_run("substantial returns").bold = True
    para.add_run(" within the first year of full deployment and position our "
                 "organization as a data-driven leader in our industry.").bold = False

    # Save document
    output_path = "sample_report.docx"
    doc.save(output_path)
    print(f"Enhanced sample document created: {output_path}")
    print("\nDocument includes:")
    print("  - Title and Subtitle")
    print("  - Heading 1, 2, and 3 levels")
    print("  - Regular paragraphs")
    print("  - Bullet lists")
    print("  - Numbered lists")
    print("  - 4 TABLES with varying column counts:")
    print("      Table 1: 2 columns (Capability | Impact)")
    print("      Table 2: 3 columns (Phase | Timeline | Deliverables)")
    print("      Table 3: 4 columns (Category | Year 1 | Year 2 | Year 3)")
    print("      Table 4: 3 columns (Risk | Impact | Mitigation)")
    print("  - Bold and italic formatting")
    print("\nReady for brand styling with professional table formatting!")


if __name__ == "__main__":
    create_sample_document()
