#!/usr/bin/env python3
"""
Create Demo Documents for Brand Styles
=======================================

This script creates sample documents styled with different brand themes
using the Document Polisher skill.

Usage:
    # From the repository root:
    python examples/scripts/create_demo_docs.py

    # With specific brands:
    python examples/scripts/create_demo_docs.py mckinsey deloitte stripe

Requirements:
    - python-docx: pip install python-docx
    - Run from repository root directory

Output:
    Creates styled documents in ./output/demo_documents/
"""

import os
import sys
from pathlib import Path

# Ensure we can import from the skills directory
REPO_ROOT = Path(__file__).parent.parent.parent
sys.path.insert(0, str(REPO_ROOT / '.claude' / 'skills' / 'document-polisher' / 'scripts'))

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Error: python-docx is required.")
    print("Install with: pip install python-docx")
    sys.exit(1)


def create_base_demo_document(output_path: str) -> str:
    """Create a base demo document with sample business content."""
    doc = Document()

    # Title
    doc.add_heading('Demo Document', 0)

    # Introduction
    doc.add_heading('Executive Summary', level=1)
    doc.add_paragraph(
        'This document demonstrates the brand styling capabilities of the Document Polisher skill. '
        'Each section showcases different formatting elements including headings, paragraphs, '
        'lists, and tables styled according to the selected brand guidelines.'
    )

    # Section 1
    doc.add_heading('Key Findings', level=1)
    doc.add_paragraph(
        'Our analysis reveals several important insights that will shape our strategic direction '
        'for the upcoming fiscal year. The following points summarize our core discoveries:'
    )

    # Bullet list
    doc.add_paragraph('Market share increased by 15% year-over-year', style='List Bullet')
    doc.add_paragraph('Customer satisfaction scores reached an all-time high of 94%', style='List Bullet')
    doc.add_paragraph('Operational efficiency improved through automation initiatives', style='List Bullet')
    doc.add_paragraph('New product launches exceeded revenue projections by 22%', style='List Bullet')

    # Section 2
    doc.add_heading('Strategic Recommendations', level=1)

    doc.add_heading('Short-Term Initiatives', level=2)
    doc.add_paragraph(
        'In the immediate term, we recommend focusing on consolidating gains in our core markets '
        'while preparing the groundwork for expansion into adjacent verticals. This balanced approach '
        'minimizes risk while maintaining growth momentum.'
    )

    doc.add_heading('Long-Term Vision', level=2)
    doc.add_paragraph(
        'Looking ahead, our vision encompasses becoming the market leader in sustainable solutions. '
        'This will require significant investment in research and development, strategic partnerships, '
        'and talent acquisition to build the capabilities necessary for transformation.'
    )

    # Section 3 with table
    doc.add_heading('Performance Metrics', level=1)
    doc.add_paragraph('The table below summarizes our key performance indicators:')

    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # Header row
    header_cells = table.rows[0].cells
    header_cells[0].text = 'Metric'
    header_cells[1].text = 'Target'
    header_cells[2].text = 'Actual'

    # Data rows
    data = [
        ('Revenue Growth', '12%', '15.3%'),
        ('Customer Retention', '85%', '91%'),
        ('Net Promoter Score', '50', '62'),
    ]

    for i, row_data in enumerate(data):
        row = table.rows[i + 1]
        for j, cell_text in enumerate(row_data):
            row.cells[j].text = cell_text

    # Conclusion
    doc.add_heading('Next Steps', level=1)
    doc.add_paragraph(
        'To capitalize on these findings, we propose the following action items:'
    )

    doc.add_paragraph('Convene a strategy session with senior leadership', style='List Number')
    doc.add_paragraph('Develop detailed implementation roadmaps for each initiative', style='List Number')
    doc.add_paragraph('Establish cross-functional working groups', style='List Number')
    doc.add_paragraph('Create a monitoring framework to track progress', style='List Number')

    doc.add_heading('Conclusion', level=2)
    doc.add_paragraph(
        'In conclusion, the opportunities before us are significant. With focused execution '
        'and continued commitment to excellence, we are well-positioned to achieve our ambitious goals. '
        'This document serves as our roadmap for the journey ahead.'
    )

    doc.save(output_path)
    print(f"Created base document: {output_path}")
    return output_path


def main():
    """Main entry point."""
    # Default brands if none specified
    default_brands = ['economist', 'mckinsey', 'deloitte', 'kpmg', 'stripe',
                      'apple', 'ibm', 'notion', 'linear', 'figma']

    # Use command line args or defaults
    brands = sys.argv[1:] if len(sys.argv) > 1 else default_brands

    # Setup paths
    output_dir = REPO_ROOT / 'output' / 'demo_documents'
    output_dir.mkdir(parents=True, exist_ok=True)

    script_path = REPO_ROOT / '.claude' / 'skills' / 'document-polisher' / 'scripts' / 'apply_brand.py'

    if not script_path.exists():
        print(f"Error: apply_brand.py not found at {script_path}")
        print("Make sure you're running from the repository root.")
        sys.exit(1)

    # Create base document
    base_doc_path = output_dir / 'demo_doc_base.docx'
    create_base_demo_document(str(base_doc_path))

    # Apply each brand style
    for brand in brands:
        output_path = output_dir / f'demo_doc_{brand}.docx'
        cmd = f'python "{script_path}" "{base_doc_path}" {brand} "{output_path}"'
        print(f"\nApplying {brand} style...")
        result = os.system(cmd)
        if result != 0:
            print(f"  Warning: Failed to apply {brand} style")

    print(f"\n{'='*60}")
    print(f"Demo documents created in: {output_dir}")
    print(f"{'='*60}")


if __name__ == '__main__':
    main()
